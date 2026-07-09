"""
Generate a verified rock-stability methods comparison as a Word document,
with per-method source citations and confidence notes.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEADER_BG = "1F3864"   # dark blue
ALT_BG = "D9E1F2"      # light blue


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, size=8.5, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_table(table):
    table.style = "Table Grid"
    table.autofit = True


doc = Document()

# landscape A4
sec = doc.sections[0]
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width
for m in (sec.left_margin, sec.right_margin):
    pass
sec.left_margin = Cm(1.5)
sec.right_margin = Cm(1.5)
sec.top_margin = Cm(1.5)
sec.bottom_margin = Cm(1.5)

# ---- Title ----
h = doc.add_paragraph()
r = h.add_run("Rock Stability Methods — Verified Comparison (with sources)")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = RGBColor.from_string(HEADER_BG)

sub = doc.add_paragraph()
rs = sub.add_run(
    "All validity limits and key parameters re-checked against the primary sources. "
    "Items confirmed against clean source text are marked (verified); items relying on "
    "standard literature are flagged. Confidence and citations are given in the notes "
    "section below each table."
)
rs.font.size = Pt(9)
rs.italic = True

scope = doc.add_paragraph()
sr = scope.add_run(
    "Scope of this thesis: the nearshore environment from the shoaling zone up to just past the "
    "breaker line. In CIRIA terms this is deep-to-shallow water, NOT very shallow water "
    "(Hs-toe < 70% Hso). The very-shallow inner-surf regime with fully broken heavy waves is "
    "outside scope, so the Van der Meer shallow-water form (Van Gent et al. 2004) is not used."
)
sr.font.size = Pt(9)
sr.bold = True

# ================= TABLE 1 =================
doc.add_paragraph().add_run("Table 1  Method comparison").bold = True

cols = ["Method", "Made for", "Required inputs", "Output",
        "Damage criterion", "Validity limits (verified)", "Fit for cable berm"]

rows = [
    ["Hudson (1953/59)",
     "Slope armour, permeable-core breakwater, regular waves",
     "Hs, KD, Δ, cotα",
     "W50 / Dn50",
     "Implicit fixed (~0–5%), no scale",
     "Slope armour, non-overtopped, permeable core (KD=4 non-breaking / KD=2 breaking; "
     "impermeable core only with KD=1); deep to shallow water but NOT very shallow "
     "(Hs-toe < 70% Hso); no strict cotα range published (KD tables ~1:1.5–1:3); no Tp; no N; no P; no current",
     "Low – first estimate only"],

    ["Van der Meer (1988/98)",
     "Slope armour, random waves, deep water",
     "Hs, Tm (sm), Δ, Dn50, cotα, P, N, Sd",
     "Dn50 / Ns at a chosen Sd",
     "Sd = Ae/Dn50² (start 2–3, failure ≈8–17)",
     "Deep water (h > 3·Hs-toe); slope armour; cotα = 1.5–6; P = 0.1–0.6; N < 7500; "
     "som = 0.01–0.06; ξm = 0.7–7; Δ = 1–2.1 (higher Δ→cotα≥2); separate plunging/surging; "
     "no current. Valid in the outer/deep part of the domain, drops out as waves approach breaking",
     "Low–med – slope, not berm/bed"],

    ["Van Rijn (2019)",
     "Horizontal/mild-slope bed protection, current ± waves",
     "uc, Hs, Tp, h, Δ, slope, α(ks), r, γstr",
     "required D50; + Nmr/Sd (Cheng 2002) for dynamic",
     "r = 0.4–1.0 → θcr = r·0.05 (static); Cheng bed-load (dynamic)",
     "D50 = 0.03–3 m; horizontal/mild slope; roughness h/(αD50) & Aw/(αD50) = 10–300; "
     "D50 is sieve size (D50 ≈ 1.2·Dn50); sparse field calibration for largest sizes",
     "High – only one with combined wave+current; static & dynamic"],

    ["Van Gent & Van der Werf (2014)",
     "Toe of breakwater armour slope, waves",
     "Hs, Tm-1,0, ht, Bt, tt, Dn50, Δ, cotα, roughness",
     "Nod or Dn50",
     "Nod (displaced stones in a Dn50-wide strip)",
     "ht/h = 0.10–0.30; h/Hs = 1.2–4.5; slopes 1:1.5 & 1:2; wave steepness sp = 0.015–0.04; "
     "waves only; non-overtopped; tested on fixed bed (no toe scour)",
     "Med – berm geometry fits, but assumes armour slope above + down-rush"],

    ["Herrera & Medina (2015)",
     "Toe berm, very shallow + steep (m=1/10), breaking",
     "hs, Hs0, L0p (Tp), Δ, Dn50 (std Bt=3Dn50, tt=2Dn50)",
     "Nod (Eq. 12) + 90% CI",
     "Nod = 0.5 start, 2.0 moderate, 4.0 failure (CIRIA 2007)",
     "−0.15 < hs/Hs0 < 1.5 (negative = emerged berm); −0.5 < hs/Dn50 < 5.01; "
     "steepness 0.008–0.08 (Herrera Table, p.10); Ns = 0.81–3.36; m = 1/10; "
     "standard berm (Bt=3Dn50, tt=2Dn50); breaking waves only",
     "Med–high for bedrock/steep breaking; fixed geometry, waves only"],

    ["Deltares HaSPro / RoBeD (2023)",
     "Loose-rock scour protection & cable-cover berms, offshore, current+waves",
     "Shields θ (→ uc, Hs, Tp, h, Δ, D50), WTOP, MTOC, Dcable, side slope",
     "berm geometry (Bberm/hberm, TOPC) + deformation class; external/interface/flexibility",
     "MOBthresh = 0.95 recommended (case-specific allowed); classes 0–3; reshaping/wash-away at θ > ~0.06",
     "reshaping regime; MOB = θ/θcr (App A); θcr Soulsby Eq A.6 (incl. 1.2 factor); ks = 2.5·D50; "
     "perpendicular wave+current assumed; offshore-calibrated",
     "High – only one explicitly for cable berms + 3 requirements"],
]

t1 = doc.add_table(rows=1, cols=len(cols))
style_table(t1)
for j, c in enumerate(cols):
    set_cell_text(t1.rows[0].cells[j], c, bold=True, size=8.5, color="FFFFFF")
    shade(t1.rows[0].cells[j], HEADER_BG)
for i, row in enumerate(rows):
    cells = t1.add_row().cells
    for j, val in enumerate(row):
        set_cell_text(cells[j], val, bold=(j == 0), size=8)
        if i % 2 == 1:
            shade(cells[j], ALT_BG)

# ================= TABLE 2 =================
doc.add_paragraph()
doc.add_paragraph().add_run("Table 2  Situation → method (verified, no changes)").bold = True

cols2 = ["Situation", "Primary method", "Supporting / check", "Why"]
rows2 = [
    ["Bedrock, steep nearshore, breaking waves", "Herrera & Medina (2015)",
     "Van Gent (range check); Hudson (first estimate)",
     "Only method for steep bottom + breaking; shows that gentle-slope formulas fail here."],
    ["Mobile sand, combined waves + current (crossing/obstacle)", "Van Rijn (2019)",
     "Deltares RoBeD (deformation + interface/flexibility)",
     "Only method with genuine wave+current on a bed; Deltares covers winnowing/falling apron on sand."],
    ["Cable berm geometry & reshaping/deformation (any)", "Deltares RoBeD (2023)",
     "Van Rijn (D50 sizing)",
     "Provides berm geometry + deformation class rather than only D50."],
    ["Quick first sizing / sanity check", "Hudson / Van der Meer", "—",
     "Classical and fast, but without wave period or current."],
]
t2 = doc.add_table(rows=1, cols=len(cols2))
style_table(t2)
for j, c in enumerate(cols2):
    set_cell_text(t2.rows[0].cells[j], c, bold=True, size=9, color="FFFFFF")
    shade(t2.rows[0].cells[j], HEADER_BG)
for i, row in enumerate(rows2):
    cells = t2.add_row().cells
    for j, val in enumerate(row):
        set_cell_text(cells[j], val, bold=(j == 1), size=8.5)
        if i % 2 == 1:
            shade(cells[j], ALT_BG)

# ================= SOURCES & NOTES =================
doc.add_paragraph()
doc.add_paragraph().add_run("Sources & verification notes").bold = True

notes = [
    ("Hudson (1953/59)",
     "CIRIA C683 Rock Manual (2007), §5.2.2.2 (Eq. 5.133–5.135) and Table 5.29 (p.579); "
     "Schiereck & Verhagen (2019), §8.3.2, Eq. 8.10.",
     "VERIFIED against CIRIA. CIRIA gives NO strict cotα validity range for Hudson (unlike Van der "
     "Meer). Derived from regular-wave tests on permeable-core structures. KD = 4 (non-breaking) / "
     "2 (breaking, per SPM 1984); impermeable core only with KD = 1. Per Table 5.29 Hudson is "
     "applicable for deep to shallow water but NOT very shallow (Hs-toe < 70% Hso). Your earlier "
     "1.5–4 is not in CIRIA; use Schiereck’s 1:1.5–1:6 or drop the upper bound."),
    ("Van der Meer (1988/98)",
     "CIRIA C683 Rock Manual (2007), Table 5.24 (p.570, deep-water validity) and Table 5.29 "
     "(fields of application); Schiereck & Verhagen (2019), §8.3.2–8.3.3.",
     "VERIFIED against CIRIA Table 5.24 (authoritative). Deep water (h > 3·Hs-toe): cotα = 1.5–6, "
     "N < 7500, som = 0.01–0.06, ξm = 0.7–7, Δ = 1–2.1 (for Δ up to 3.5 → cotα ≥ 2), P = 0.1–0.6 "
     "(P = 0.1 impermeable, 0.4 armour+filter on permeable core, 0.5 armour on permeable core, "
     "0.6 homogeneous). Plunging/surging split. The shallow-water form (Van Gent et al. 2004, "
     "needs H2%) covers very shallow water only and is OUT OF SCOPE for this thesis."),
    ("Van Rijn (2019)",
     "van Rijn (2019), Critical movement of large rocks in currents and waves, "
     "Eq. 3 (θcr), Eq. 5 (wave+current), Table 1 (fc/fw ranges); Cheng (2002) for the dynamic case.",
     "VERIFIED against the paper. Added vs your table: roughness ratio h/(αD50) & Aw/(αD50) = 10–300, "
     "and D50 is the sieve diameter with D50 ≈ 1.2·Dn50 (convert if inputs are nominal Dn50)."),
    ("Van Gent & Van der Werf (2014)",
     "van Gent & van der Werf (2014), Rock toe stability of rubble mound breakwaters, "
     "validity statements in the paper (ht/h and h/Hs ranges, tested steepnesses).",
     "VERIFIED against the paper. Added vs your table: wave steepness sp = 0.015–0.04 (tested), "
     "non-overtopped structure, and the tests used a fixed (non-mobile) bed so toe scour was excluded. "
     "Uses Tm-1,0 (spectral period) and Hs = Hm0."),
    ("Herrera & Medina (2015)",
     "Herrera & Medina (2015), Toe berm design for very shallow waters on steep sea bottoms, "
     "Eq. 12; validity ranges §2; damage levels after CIRIA Rock Manual (2007).",
     "VERIFIED against the paper (pages you sent). −0.15 < hs/Hs0 < 1.5 is correct (negative lower "
     "bound = emerged berm). Damage levels per CIRIA (2007): Nod = 0.5 (start), 2.0 (moderate), "
     "4.0 (failure). REMOVED “Nod = 1 = design” – it is not in Herrera’s cited criteria. "
     "CORRECTED steepness: Herrera’s own validity table (p.10, column “This study”) gives 0.008–0.08, "
     "not 0.02–0.07, and Ns = 0.81–3.36 (verify the exact steepness symbol s0p vs stp when you cite it). "
     "SCOPE NOTE: Herrera’s domain (very shallow, m=1/10, heavy breaking) sits at the edge of or just "
     "outside this thesis’s nearshore domain, so it will rarely be the governing method here."),
    ("Deltares HaSPro / RoBeD (2023)",
     "Deltares (2023), Handbook of Scour and Cable Protection Methods, §5.2–5.3 (RoBeD, p.73–75), "
     "Appendix A (mobility number, Eq. A.1 & A.6). Polynomial coefficients themselves are in the "
     "loose-rock analysis report Deltares (2023n), which is not in our set.",
     "The handbook’s RoBeD section explicitly recommends MOBthresh = 0.95 for the rock-berm model, "
     "so “no fixed default” understates the source. MOB = θ/θcr; θcr from Soulsby (Eq. A.6, which "
     "keeps the 1.2 factor) and ks = 2.5·D50 for the mobility calc – note these differ from the "
     "Van Rijn (2019) θcr form used elsewhere."),
]
for name, src, note in notes:
    p = doc.add_paragraph()
    p.add_run(name + " — ").bold = True
    p.add_run("Source: ").italic = True
    p.add_run(src)
    p2 = doc.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.6)
    rr = p2.add_run(note)
    rr.font.size = Pt(9)

out = "/tmp/claude-0/-home-user-Willempol87/292b0f39-8fbd-59c1-9ef9-38af8e880047/scratchpad/Rock_Stability_Methods_Verified.docx"
doc.save(out)
print("saved:", out)
